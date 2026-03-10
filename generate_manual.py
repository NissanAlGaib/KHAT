"""
Generate PawLink User Manual in DOCX and PDF formats.
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from fpdf import FPDF

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "docs")

# ─── DOCX GENERATION ──────────────────────────────────────────────

def create_docx():
    doc = Document()
    
    # -- Style setup --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    style_h1 = doc.styles["Heading 1"]
    style_h1.font.size = Pt(22)
    style_h1.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    style_h1.font.bold = True
    
    style_h2 = doc.styles["Heading 2"]
    style_h2.font.size = Pt(16)
    style_h2.font.color.rgb = RGBColor(0x2D, 0x6B, 0xCE)
    style_h2.font.bold = True
    
    style_h3 = doc.styles["Heading 3"]
    style_h3.font.size = Pt(13)
    style_h3.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    style_h3.font.bold = True

    style_h4 = doc.styles["Heading 4"]
    style_h4.font.size = Pt(11)
    style_h4.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    style_h4.font.bold = True
    
    # -- Page setup --
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph("")
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PawLink")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Mobile App User Manual")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph("")
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 1.5.0  •  March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Platform: Android (Expo / EAS)")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    doc.add_page_break()
    
    # ── Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Getting Started",
        "3. Identity Verification",
        "4. Navigating the App",
        "5. Managing Your Pets",
        "6. Finding Breeding Matches",
        "7. Match Requests",
        "8. Chat & Messaging",
        "9. Breeding Contracts",
        "10. Shooter Workflow",
        "11. Subscriptions & Payments",
        "12. Activity & Notifications",
        "13. Profile & Settings",
        "14. Troubleshooting & FAQ",
        "15. Contact & Support",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ── Helper functions ──
    def add_table(doc, headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        # Rows
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph("")
    
    def add_note(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("📌 Note: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        run2 = p.add_run(text)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    def add_important(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("⚠️ Important: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        run2 = p.add_run(text)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    def add_bullet(doc, text, level=0):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    
    def add_numbered(doc, text, level=0):
        p = doc.add_paragraph(text, style="List Number")
        p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 1: Introduction
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_paragraph(
        "PawLink is a mobile pet breeding matchmaking platform that connects pet breeders with "
        "compatible breeding partners for their dogs and cats. The app also supports Shooters — "
        "breeding handlers/facilitators who assist with the logistics of breeding arrangements."
    )
    
    doc.add_heading("Key Features", level=3)
    features = [
        ("Smart Matching", "Swipe-based interface to discover compatible breeding partners with compatibility scores."),
        ("AI-Powered Tools", "Breed identification from photos, offspring appearance prediction, and OCR document scanning."),
        ("Breeding Contracts", "In-app contract creation and management for breeding arrangements."),
        ("Secure Payments", "Integrated PayMongo payment system with money pool escrow for collateral and compensation."),
        ("Real-Time Chat", "Direct messaging between matched breeders with read receipts."),
        ("Verification System", "Identity and document verification to ensure trust and safety."),
        ("Vaccination Tracking", "Comprehensive health record management for your pets."),
    ]
    for name, desc in features:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name} — ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 2: Getting Started
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Getting Started", level=1)
    
    doc.add_heading("2.1 Installing the App", level=2)
    add_numbered(doc, "Download the PawLink APK from the provided distribution link or scan the QR code shared by the development team.")
    add_numbered(doc, 'On your Android device, go to Settings > Security and enable "Install from unknown sources" (if not already enabled).')
    add_numbered(doc, "Open the downloaded APK file and tap Install.")
    add_numbered(doc, "Once installed, tap Open to launch PawLink.")
    add_note(doc, "PawLink receives over-the-air (OTA) updates automatically. When an update is available, you will be prompted to restart the app.")
    
    doc.add_heading("2.2 Creating an Account", level=2)
    doc.add_paragraph("Tap Register on the login screen to begin the 4-step registration process:")
    
    doc.add_heading("Step 1 — Account Setup", level=4)
    add_bullet(doc, "Enter your email address (must be valid and unique).")
    add_bullet(doc, "Choose a username.")
    add_bullet(doc, "Create a password with the following requirements:")
    add_bullet(doc, "At least 8 characters long", 1)
    add_bullet(doc, "At least one uppercase letter", 1)
    add_bullet(doc, "At least one number", 1)
    add_bullet(doc, "At least one special character", 1)
    
    doc.add_heading("Step 2 — Personal Information", level=4)
    add_bullet(doc, "Enter your first name and last name.")
    add_bullet(doc, "Select your birthdate (you must be at least 13 years old).")
    add_bullet(doc, "Choose your sex (Male or Female).")
    
    doc.add_heading("Step 3 — Address", level=4)
    doc.add_paragraph("Enter your address in the Philippine format:")
    add_bullet(doc, "Street Address")
    add_bullet(doc, "Barangay")
    add_bullet(doc, "City / Municipality")
    add_bullet(doc, "Province")
    add_bullet(doc, "Postal Code (4 digits)")
    
    doc.add_heading("Step 4 — Role Selection", level=4)
    doc.add_paragraph("Choose one or both roles:")
    add_bullet(doc, "Breeder — You own pets and are looking for breeding partners.")
    add_bullet(doc, "Shooter — You facilitate and handle breeding arrangements.")
    doc.add_paragraph("You must select at least one role. You can select both.")
    doc.add_paragraph("Tap Submit to create your account. You will be redirected to the login screen.")
    
    doc.add_heading("2.3 Logging In", level=2)
    add_numbered(doc, "Enter your email and password.")
    add_numbered(doc, "Tap Sign In.")
    add_numbered(doc, "You will be taken to the Home screen.")
    add_note(doc, "If your account has been suspended or banned, an alert will appear with an explanation.")
    
    doc.add_heading("2.4 Forgot Password", level=2)
    doc.add_paragraph("If you forget your password, tap Forgot Password? on the login screen. Follow the on-screen instructions to reset your password via email.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 3: Identity Verification
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Identity Verification", level=1)
    doc.add_paragraph("Before you can add pets or send match requests, you must verify your identity.")
    
    doc.add_heading("3.1 Verifying Your Identity", level=2)
    add_numbered(doc, 'From the Profile tab, tap on your My Pets section and then the "+" button, or navigate to Settings > Verification Status.')
    add_numbered(doc, "If you have not yet verified, you will be guided through the verification wizard.")
    
    doc.add_heading("Step 1 — Government ID Verification (Required)", level=4)
    add_bullet(doc, "Tap Upload ID to take a photo or select one from your gallery.")
    add_bullet(doc, "The app uses OCR (Optical Character Recognition) to automatically read and fill in details from your ID.")
    add_bullet(doc, "Fields auto-filled: Full Name, ID Number, Issuing Authority, Issue and Expiry Dates.")
    add_bullet(doc, "Review the auto-filled information and make corrections if needed.")
    add_bullet(doc, "Tap Next to proceed.")
    
    doc.add_heading("Step 2 — Licensed Breeder Certificate (Optional)", level=4)
    add_bullet(doc, "If you have a breeder's license, upload it here.")
    add_bullet(doc, "OCR will auto-fill certificate details.")
    add_bullet(doc, "You may tap Skip if you don't have one.")
    
    doc.add_heading("Step 3 — Shooter Certificate (Optional)", level=4)
    add_bullet(doc, "If you are a Shooter with certification, upload it here.")
    add_bullet(doc, "You may tap Skip if not applicable.")
    doc.add_paragraph("Tap Submit to send your documents for review.")
    
    doc.add_heading("3.2 Adding Certificates Later", level=2)
    add_numbered(doc, "Go to Settings > Verification Status.")
    add_numbered(doc, "Tap Add Certificate on the relevant document type.")
    add_numbered(doc, "Upload and submit.")
    
    doc.add_heading("3.3 Checking Verification Status", level=2)
    doc.add_paragraph("Go to Settings > Verification Status to see the current state of each document:")
    add_table(doc,
        ["Status", "Meaning"],
        [
            ["Not Submitted", "You have not uploaded this document yet."],
            ["Under Review", "Your document has been submitted and is awaiting admin review."],
            ["Verified ✅", "Your document has been approved."],
            ["Rejected ❌", "Your document was rejected. Check the reason and resubmit."],
        ]
    )
    doc.add_paragraph("If a document is rejected, tap Resubmit to upload a new version with corrections.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 4: Navigating the App
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Navigating the App", level=1)
    
    doc.add_heading("4.1 Bottom Navigation Tabs", level=2)
    doc.add_paragraph("PawLink uses a bottom navigation bar with 5 tabs:")
    add_table(doc,
        ["Tab", "Description"],
        [
            ["Home", "Your main dashboard — browse top matches, nearby pets, and shooters."],
            ["Matches", "Manage match requests, view active matches, and review history."],
            ["Match (center)", "Full-screen swipe interface for discovering pet matches."],
            ["Activity", "Notification center — verification, match alerts, messages, system notices."],
            ["Profile", "Your profile dashboard, pet list, and settings."],
        ]
    )
    
    doc.add_heading("4.2 Switching Roles (Breeder / Shooter)", level=2)
    doc.add_paragraph("If you registered for both Breeder and Shooter roles:")
    add_numbered(doc, "Go to the Profile tab.")
    add_numbered(doc, "Tap the Role Switcher toggle near the top of the screen.")
    add_numbered(doc, "The app will switch between Breeder and Shooter modes.")
    add_bullet(doc, "Breeder mode: Home shows the matching interface and pet discovery.")
    add_bullet(doc, "Shooter mode: Home shows the Shooter dashboard with offers and assignments.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 5: Managing Your Pets
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. Managing Your Pets", level=1)
    
    doc.add_heading("5.1 Adding a New Pet", level=2)
    add_important(doc, "You must complete identity verification before adding a pet.")
    add_numbered(doc, "Go to the Profile tab.")
    add_numbered(doc, "Tap the My Pets section.")
    add_numbered(doc, 'Tap the "+" floating action button.')
    add_numbered(doc, "Complete the 5-step Add Pet wizard:")
    
    doc.add_heading("Step 1 — Pet Information", level=4)
    add_bullet(doc, "Enter your pet's name.")
    add_bullet(doc, "Select species (Dog or Cat).")
    add_bullet(doc, "AI Breed Identification: Tap the camera icon to take a photo or select from gallery. The AI identifies the breed and auto-fills it.")
    add_bullet(doc, "Alternatively, manually select the breed from the dropdown.")
    add_bullet(doc, "Select sex (Male or Female).")
    add_bullet(doc, "Enter birthdate, height (cm), and weight (kg).")
    add_bullet(doc, "Optionally enter a microchip ID.")
    
    doc.add_heading("Step 2 — About Your Pet", level=4)
    add_bullet(doc, "Select relevant behaviors from the multi-select list (e.g., Friendly, Playful, Calm, Protective).")
    add_bullet(doc, "Select relevant attributes (e.g., Hypoallergenic, Good with children).")
    add_bullet(doc, "Write a description about your pet.")
    
    doc.add_heading("Step 3 — Health Certificate", level=4)
    add_bullet(doc, "Upload your pet's health certificate document (photo).")
    add_bullet(doc, "OCR will auto-fill certificate details.")
    
    doc.add_heading("Step 4 — Photos", level=4)
    add_bullet(doc, "Upload a minimum of 3 photos of your pet.")
    add_bullet(doc, "Tap on a photo to set it as the primary photo (shown in match cards).")
    
    doc.add_heading("Step 5 — Preferences", level=4)
    add_bullet(doc, "Set your preferred breeding partner criteria: breed, behaviors, attributes, age range.")
    doc.add_paragraph("Tap Submit to register your pet. After success, you may import vaccination history.")
    
    doc.add_heading("5.2 Viewing Your Pet's Profile", level=2)
    doc.add_paragraph("From Profile > My Pets, tap any pet card. The profile has 4 tabs:")
    add_bullet(doc, "About — Species, breed, sex, age, weight, height, microchip ID, behaviors, attributes, and description.")
    add_bullet(doc, "Health — Vaccination cards with shot records and document status (Valid ✅ / Expiring ⚠️ / Expired ❌).")
    add_bullet(doc, "Gallery — Photo grid with primary photo highlighted.")
    add_bullet(doc, "History — Breeding and litter history.")
    
    doc.add_paragraph("Pet Status Badges:")
    add_table(doc,
        ["Badge", "Meaning"],
        [
            ["Available", "Your pet is ready for matching."],
            ["Pending", "Active match request or awaiting verification."],
            ["Disabled", "Temporarily disabled."],
            ["Archived", "Not active."],
            ["Cooldown: Xd", "Post-breeding cooldown period (X days remaining)."],
        ]
    )
    
    doc.add_heading("5.3 Managing Vaccinations", level=2)
    add_numbered(doc, "Open your pet's profile and go to the Health tab.")
    add_numbered(doc, "View the vaccination dashboard: Total Cards, Verified, Pending, Overdue.")
    add_numbered(doc, "To add a new shot: Tap vaccine card > Add Shot > enter date and upload photo proof > Submit.")
    add_numbered(doc, "Manage protocols: Tap Manage Protocol to opt-in to new or change existing protocols.")
    
    doc.add_heading("5.4 Importing Vaccination History", level=2)
    doc.add_paragraph("After registering a new pet, you can import past vaccination records:")
    add_numbered(doc, "When prompted after pet registration, tap Import Vaccination History.")
    add_numbered(doc, "For each vaccine card, enter shot dates and upload supporting documents.")
    add_numbered(doc, "Submit all records.")
    
    doc.add_heading("5.5 Viewing Litters & Breeding History", level=2)
    add_numbered(doc, "Open your pet's profile History tab or navigate to the Litters screen.")
    add_numbered(doc, "View summary statistics: Total Litters, Total Pups/Kittens, Alive Count.")
    add_numbered(doc, "Tap a litter card for details: Overview, Offspring (with allocation status), Health.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 6: Finding Breeding Matches
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Finding Breeding Matches", level=1)
    
    doc.add_heading("6.1 The Home Screen (Breeder)", level=2)
    doc.add_paragraph("In Breeder mode, the Home screen displays:")
    add_bullet(doc, "Greeting Header — Personalized greeting with search icon and subscription badge.")
    add_bullet(doc, "Top Matches — A swipeable card stack showing best-matched pets with photo, name, breed, and compatibility score.")
    add_bullet(doc, "Nearby Pets — Horizontally scrollable list of pets in your area.")
    add_bullet(doc, "Shooters — Horizontally scrollable list of available shooters.")
    add_bullet(doc, 'Browse Tabs — Toggle between full "Pets" and "Shooters" lists.')
    
    doc.add_heading("6.2 Swiping Through Matches", level=2)
    doc.add_paragraph("On the Home screen and the Match tab (center button):")
    add_bullet(doc, "Swipe Right or tap ♥ (Heart) to send a match request.")
    add_bullet(doc, "Swipe Left or tap ✕ (X) to pass. Passed pets will not be shown again.")
    add_important(doc, "Only pets of the same species and opposite sex are shown. Your own pets are excluded.")
    
    doc.add_heading("6.3 Viewing a Pet's Public Profile", level=2)
    doc.add_paragraph("Tap any pet card to view the public profile with 5 tabs:")
    add_bullet(doc, "About — Species, breed, age, weight, height, description, behavior/attribute tags.")
    add_bullet(doc, "Health — Vaccination status and health certification.")
    add_bullet(doc, "Gallery — Photo gallery.")
    add_bullet(doc, "Litters — Previous breeding/litter history.")
    add_bullet(doc, "Compatibility — Detailed score between this pet and your selected pet, with trait-by-trait breakdown.")
    doc.add_paragraph('A "Send Match Request" button appears at the bottom.')
    
    doc.add_heading("6.4 Checking Compatibility", level=2)
    doc.add_paragraph("On a pet's public profile, go to the Compatibility tab to see:")
    add_bullet(doc, "Overall compatibility percentage score.")
    add_bullet(doc, "Trait-by-trait breakdown: breed compatibility, health status, age, location, behavior matching, and preferences.")
    
    doc.add_heading("6.5 AI Offspring Prediction", level=2)
    add_numbered(doc, "From a pet's public profile, tap AI Offspring Prediction.")
    add_numbered(doc, "View both parent pets side by side with compatibility score.")
    add_numbered(doc, "Tap Generate Prediction to create an AI-generated offspring image.")
    add_numbered(doc, "View predicted traits: breed mix, coat type, size estimate, temperament.")
    add_note(doc, "AI generations per day are limited by your subscription tier (1/day for Free plan).")
    
    doc.add_heading("6.6 Searching for Pets, Breeders & Shooters", level=2)
    add_numbered(doc, "Tap the search icon on the Home screen.")
    add_numbered(doc, "Type your query (results appear after a brief delay).")
    add_numbered(doc, "Filter by category: All, Pets, Breeders, or Shooters.")
    add_numbered(doc, "For Pets, further filter by Species (Dog/Cat) and Sex (Male/Female).")
    add_numbered(doc, "Tap a result to navigate to the profile.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 7: Match Requests
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Match Requests", level=1)
    
    doc.add_heading("7.1 Sending a Match Request", level=2)
    doc.add_paragraph("Ways to send a match request:")
    add_bullet(doc, "Swipe right on a match card.")
    add_bullet(doc, "Tap the ♥ button on a match card.")
    add_bullet(doc, 'Tap "Send Match Request" on a pet\'s public profile.')
    doc.add_paragraph("Requirements: Identity must be verified, active pet selected. Free tier may require per-request payment.")
    
    doc.add_heading("7.2 Managing Incoming Requests", level=2)
    add_numbered(doc, "Go to the Matches tab > Requests > Incoming.")
    add_numbered(doc, "Each card shows: pet photo, name, breed, distance, compatibility score.")
    add_numbered(doc, "Tap ♥ to accept — a conversation is created automatically.")
    add_numbered(doc, "Tap ✕ to decline.")
    
    doc.add_heading("7.3 Managing Outgoing Requests", level=2)
    add_numbered(doc, "Go to Matches tab > Requests > Outgoing.")
    add_numbered(doc, "Tap Cancel to withdraw a pending request.")
    
    doc.add_heading("7.4 Match History", level=2)
    add_numbered(doc, "In the Matches tab, go to the History section.")
    add_numbered(doc, "Filter by: All, Declined, or Cancelled.")
    add_numbered(doc, "Browse previous match requests with pagination.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 8: Chat & Messaging
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. Chat & Messaging", level=1)
    
    doc.add_heading("8.1 Starting a Conversation", level=2)
    doc.add_paragraph('After a match request is accepted, tap "Start Chat" on the matched pair in the Matches tab, or tap the notification in the Activity tab.')
    
    doc.add_heading("8.2 Sending Messages", level=2)
    doc.add_paragraph("The conversation screen provides:")
    add_bullet(doc, "Real-time messaging — Messages refresh every 5 seconds.")
    add_bullet(doc, "Message bubbles — Your messages on the right (blue), theirs on the left.")
    add_bullet(doc, "Timestamps on each message.")
    add_bullet(doc, "Read receipts — Double check marks (✓✓) when the other party has read your message.")
    add_bullet(doc, "Date headers between message groups.")
    add_bullet(doc, "Match information at the top: matched pet, other breeder, and shooter (if assigned).")
    add_bullet(doc, "Contract card — If a contract exists, a compact card appears. Tap to view details.")
    add_bullet(doc, "Match timeline showing breeding match progress.")
    
    doc.add_heading("8.3 Blocking & Reporting Users", level=2)
    add_numbered(doc, "Tap the shield icon in the chat header.")
    add_numbered(doc, "Choose: Block User (prevents further communication) or Report User (select a reason and submit).")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 9: Breeding Contracts
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("9. Breeding Contracts", level=1)
    
    doc.add_heading("9.1 Creating a Contract", level=2)
    doc.add_paragraph('In the conversation screen, tap "Create Contract" (appears when no contract exists). Complete 5 steps:')
    
    doc.add_heading("Step 1 — Compensation Type", level=4)
    add_bullet(doc, "Money (₱) — Monetary payment.")
    add_bullet(doc, "Share Offspring — Share puppies/kittens from the litter.")
    add_bullet(doc, "Goods / Foods — Non-monetary compensation.")
    doc.add_paragraph("You can combine multiple types.")
    
    doc.add_heading("Step 2 — Compensation Details", level=4)
    add_bullet(doc, "Money: Enter the payment amount (₱).")
    add_bullet(doc, "Offspring Sharing: Choose split method (Percentage / Fixed / Alternating), enter value, select who picks first.")
    add_bullet(doc, "Goods/Foods: Describe the items.")
    
    doc.add_heading("Step 3 — Shooter (Optional)", level=4)
    add_bullet(doc, "Search for or enter a shooter's name.")
    add_bullet(doc, "Enter shooter payment (₱), location description, and special conditions.")
    
    doc.add_heading("Step 4 — Collateral & Timeline", level=4)
    add_bullet(doc, "Security Deposit: Total amount, split 50/50 between both parties, held in escrow.")
    add_bullet(doc, "Contract End Date: Set the deadline.")
    
    doc.add_heading("Step 5 — Review & Submit", level=4)
    add_bullet(doc, "Review the full summary and tap Submit.")
    
    doc.add_heading("9.2 Reviewing & Accepting a Contract", level=2)
    add_numbered(doc, "Open the conversation and tap the contract card.")
    add_numbered(doc, "Review all terms on the contract detail screen.")
    add_numbered(doc, "Tap Accept to agree, or Reject to decline.")
    add_numbered(doc, "If accepted, both parties pay their share of collateral.")
    
    doc.add_heading("9.3 Contract Lifecycle", level=2)
    add_table(doc,
        ["Stage", "Description"],
        [
            ["1. Create Contract", "One party drafts the contract terms."],
            ["2. Accept Contract", "The other party reviews and accepts (or rejects)."],
            ["3. Pay Collateral", "Both parties pay 50% each via PayMongo. Funds held in pool."],
            ["4. Submit Reports", "Both parties submit breeding progress reports."],
            ["5. Mark Breeding", "When breeding occurs, mark as complete."],
            ["6. Record Offspring", "Record litter — number of offspring, health status."],
            ["7. Complete Match", "Finalize contract. Funds released from pool."],
        ]
    )
    
    doc.add_paragraph("Contract Statuses:")
    add_table(doc,
        ["Status", "Meaning"],
        [
            ["Draft", "Contract is being created."],
            ["Pending", "Waiting for the other party to review."],
            ["Accepted", "Both parties agree — awaiting collateral payment."],
            ["Shooter Requested", "A shooter has been invited to the contract."],
            ["Shooter Accepted", "The shooter has accepted the assignment."],
            ["In Progress", "Breeding activities are underway."],
            ["Breeding Complete", "Breeding has been confirmed."],
            ["Completed", "Contract is finalized and closed."],
            ["Rejected", "Contract was rejected by one party."],
            ["Cancelled", "Contract was cancelled."],
        ]
    )
    
    doc.add_heading("9.4 Contract Detail Tabs", level=2)
    add_bullet(doc, 'Overview — Full contract terms, compensation, timeline, status. Animated "Next Action" banner guides you.')
    add_bullet(doc, "Payments — Track collateral, compensation, and shooter payments.")
    add_bullet(doc, "Reports — Submit and view breeding progress reports.")
    add_bullet(doc, "Breeding — Log breeding events and mark completion.")
    add_bullet(doc, "Offspring — Record litter details, individual offspring, and allocation.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 10: Shooter Workflow
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("10. Shooter Workflow", level=1)
    
    doc.add_heading("10.1 Shooter Home Dashboard", level=2)
    doc.add_paragraph("Switch to Shooter role to access the Shooter Home:")
    add_bullet(doc, "Stats Banner: Active, Pending, Completed, Available counts + pet breakdown (dogs/cats).")
    add_bullet(doc, "Current tab: Active and pending assignments with both pets, payment, location, breeding status.")
    add_bullet(doc, "Available tab: Open offers that can be accepted.")
    add_bullet(doc, "Finished tab: Completed or failed assignments with outcome badges.")
    
    doc.add_heading("10.2 Browsing Available Offers", level=2)
    add_numbered(doc, "In Shooter Home, tap the Available tab.")
    add_numbered(doc, "Browse offers showing both pets, payment amount (₱), and location.")
    add_numbered(doc, "Tap an offer to view details.")
    
    doc.add_heading("10.3 Accepting a Shooter Offer", level=2)
    add_numbered(doc, "On the Offer Details screen, review both pets and payment.")
    add_numbered(doc, 'Tap "Accept Offer" — breeders are notified for confirmation.')
    
    doc.add_heading("10.4 Active Assignments", level=2)
    add_numbered(doc, "Active assignments appear under the Current tab.")
    add_numbered(doc, "Tap to join the conversation between breeders.")
    add_numbered(doc, "Assist with breeding logistics, submit reports, and coordinate.")
    add_numbered(doc, "Payment is processed through the money pool upon completion.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 11: Subscriptions & Payments
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("11. Subscriptions & Payments", level=1)
    
    doc.add_heading("11.1 Subscription Plans", level=2)
    add_table(doc,
        ["Feature", "Free", "Standard", "Premium"],
        [
            ["Registered Pets", "1", "More", "Unlimited"],
            ["Matches per Month", "3", "More", "Unlimited"],
            ["AI Predictions/Day", "1", "More", "Unlimited"],
            ["Billing", "—", "Monthly / Yearly", "Monthly / Yearly"],
        ]
    )
    
    doc.add_heading("11.2 Upgrading Your Plan", level=2)
    add_numbered(doc, "Go to Profile > Settings > Subscription, or tap the subscription badge on Home.")
    add_numbered(doc, "Browse plans as carousel cards. Toggle Monthly/Yearly billing.")
    add_numbered(doc, "Tap Subscribe on your chosen plan.")
    add_numbered(doc, "Complete payment via PayMongo.")
    add_numbered(doc, "Success → redirected with confirmation, plan active immediately. Cancelled → cancellation notice.")
    
    doc.add_heading("11.3 My Payments & Money Pool", level=2)
    doc.add_paragraph("Access via Profile > Settings > My Payments.")
    doc.add_paragraph("The Money Pool is a virtual escrow for breeding contract finances:")
    add_table(doc,
        ["Type", "Description"],
        [
            ["Held", "Funds in escrow for active contracts."],
            ["Frozen", "Funds frozen due to dispute or admin action."],
            ["Pending Deposits", "Collateral payments still processing."],
            ["Total Released", "Funds released from completed contracts."],
        ]
    )
    doc.add_paragraph("Transaction Tabs: Deposits, Releases, Refunds, Fee Deductions.")
    doc.add_paragraph("Color coding: Green = earnings, Orange = credits/pending, Red = debits/deductions.")
    
    doc.add_heading("11.4 Filing a Dispute", level=2)
    add_numbered(doc, "Navigate to My Payments > Disputes.")
    add_numbered(doc, "Tap File Dispute and provide details.")
    add_numbered(doc, "Submit for admin review and track status in the Disputes tab.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 12: Activity & Notifications
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("12. Activity & Notifications", level=1)
    
    doc.add_paragraph("The Activity tab serves as your unified notification center.")
    doc.add_paragraph("Filter notifications by category:")
    add_bullet(doc, "All — View all notifications.")
    add_bullet(doc, "Verification — Document verification updates.")
    add_bullet(doc, "Matches — Match request updates.")
    add_bullet(doc, "Messages — New chat messages.")
    add_bullet(doc, "Shooter — Shooter assignments and offers.")
    add_bullet(doc, "Subscription — Plan changes and payment confirmations.")
    add_bullet(doc, "System — General announcements and alerts.")
    
    doc.add_paragraph("Verification notifications are grouped by severity:")
    add_bullet(doc, "⚠️ Warnings — Documents expiring soon.")
    add_bullet(doc, "🔴 Rejected / Needs Attention — Rejected with reasons; tap Resubmit.")
    add_bullet(doc, "🟡 Under Review — Submitted and pending.")
    add_bullet(doc, "🟢 Approved — Verified and valid.")
    
    doc.add_paragraph('Tap "Mark All as Read" to clear unread indicators. Tap any notification to navigate to the relevant screen.')
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 13: Profile & Settings
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("13. Profile & Settings", level=1)
    
    doc.add_heading("13.1 Viewing Your Profile", level=2)
    doc.add_paragraph("The Profile tab shows:")
    add_bullet(doc, "Profile header — Photo, name, verification badge.")
    add_bullet(doc, "Role switcher — Toggle Breeder/Shooter.")
    add_bullet(doc, "Dashboard tab — Stats: Current Breeding, Total Matches, Success Rate, Income (₱).")
    add_bullet(doc, 'My Pets tab — Grid of pets with status badges. "+" button to add new pet.')
    add_bullet(doc, "Settings tab — Quick access to account settings.")
    
    doc.add_heading("13.2 Editing Your Profile", level=2)
    add_numbered(doc, "Go to Profile > Settings > Account (or tap your profile header).")
    add_numbered(doc, "Update: First Name, Last Name, Contact Number, Birthdate, Sex, Profile Image.")
    add_numbered(doc, "Email is displayed but cannot be changed.")
    add_numbered(doc, "Tap Save to apply changes.")
    
    doc.add_heading("13.3 Privacy & Security", level=2)
    doc.add_paragraph("Go to Profile > Settings > Privacy & Security:")
    
    doc.add_heading("Change Password", level=4)
    add_numbered(doc, "Enter current password.")
    add_numbered(doc, "Enter new password (8+ chars, uppercase, number, special char).")
    add_numbered(doc, "Confirm and tap Update Password.")
    
    doc.add_heading("Delete Account", level=4)
    add_numbered(doc, "Tap Delete Account.")
    add_numbered(doc, 'Type "DELETE" in the confirmation field.')
    add_numbered(doc, "Enter your password.")
    add_numbered(doc, "Confirm — This is permanent and cannot be undone.")
    
    doc.add_heading("13.4 Signing Out", level=2)
    add_numbered(doc, "Go to Profile > Settings.")
    add_numbered(doc, "Tap Sign Out.")
    add_numbered(doc, "Confirm to log out.")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 14: Troubleshooting
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("14. Troubleshooting & FAQ", level=1)
    
    faqs = [
        ("Q: I can't add a pet. What do I do?",
         "A: You must complete identity verification first. Go to Settings > Verification Status and ensure your government ID is verified."),
        ("Q: My verification document was rejected.",
         "A: Go to Settings > Verification Status, check the rejection reason, and tap Resubmit."),
        ("Q: I can't see any pets to match with.",
         "A: Make sure you have a pet registered and selected. Matches only show same species, opposite sex. Try the Search feature."),
        ("Q: I was charged but my subscription didn't activate.",
         "A: Return to the app and check Settings > Subscription. If the issue persists, contact support."),
        ("Q: What does 'Pending Shooter Request' mean?",
         "A: A shooter has been invited but hasn't accepted yet. Wait for their response."),
        ("Q: I accidentally passed on a pet. Can I undo it?",
         "A: Passed pets are not shown again in the swipe view. Use Search to find and send a match request from their profile."),
        ("Q: My account is suspended or banned.",
         "A: A banned screen shows the reason and end date. Contact support at pawlink.support@gmail.com."),
        ("Q: How do I switch between Breeder and Shooter?",
         "A: Go to the Profile tab and tap the role switcher toggle at the top."),
        ("Q: How does the money pool work?",
         "A: It's a virtual escrow that holds collateral and compensation during contracts. Funds are released on completion or refunded on cancellation."),
        ("Q: How many pets can I register?",
         "A: Free plan: 1 pet. Paid plans allow more. Check the Subscription screen for details."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.bold = True
        doc.add_paragraph(a)
        doc.add_paragraph("")
    
    doc.add_page_break()
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 15: Contact & Support
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("15. Contact & Support", level=1)
    doc.add_paragraph("If you need assistance, have questions, or want to report an issue:")
    add_bullet(doc, "Email: pawlink.support@gmail.com")
    add_bullet(doc, "In-App: Use the Report feature in any conversation to flag concerns.")
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PawLink — Connecting Responsible Pet Breeders")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run("© 2026 KHAT Development Team. All rights reserved.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    # Save
    path = os.path.join(OUTPUT_DIR, "PawLink_User_Manual.docx")
    doc.save(path)
    print(f"DOCX saved: {path}")
    return path

# ─── PDF GENERATION ──────────────────────────────────────────────

class PawLinkPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=20)
    
    def header(self):
        if self.page_no() > 1:
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(148, 163, 184)
            self.cell(0, 10, "PawLink User Manual", align="L")
            self.cell(0, 10, f"Page {self.page_no()}", align="R")
            self.ln(15)
    
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, "© 2026 KHAT Development Team", align="C")
    
    def section_title(self, text):
        self.set_font("Helvetica", "B", 18)
        self.set_text_color(26, 86, 219)
        self.cell(0, 12, text, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(26, 86, 219)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(6)
    
    def subsection_title(self, text):
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(45, 107, 206)
        self.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
    
    def sub3_title(self, text):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(59, 130, 246)
        self.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)
    
    def sub4_title(self, text):
        self.set_font("Helvetica", "BI", 10)
        self.set_text_color(75, 85, 99)
        self.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)
    
    def body_text(self, text):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(51, 51, 51)
        self.multi_cell(0, 5.5, text)
        self.ln(2)
    
    def bullet(self, text, indent=10):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(51, 51, 51)
        x = self.get_x()
        self.set_x(x + indent)
        self.cell(5, 5.5, "-")
        self.multi_cell(0, 5.5, text)
        self.ln(1)
    
    def numbered_item(self, num, text, indent=10):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(51, 51, 51)
        x = self.get_x()
        self.set_x(x + indent)
        self.cell(8, 5.5, f"{num}.")
        self.multi_cell(0, 5.5, text)
        self.ln(1)
    
    def note_box(self, text):
        self.set_font("Helvetica", "I", 9)
        self.set_text_color(100, 116, 139)
        self.set_fill_color(241, 245, 249)
        x = self.get_x()
        self.set_x(x + 10)
        self.multi_cell(0, 5, f"Note: {text}", fill=True)
        self.ln(3)
    
    def important_box(self, text):
        self.set_font("Helvetica", "BI", 9)
        self.set_text_color(220, 38, 38)
        x = self.get_x()
        self.set_x(x + 10)
        self.multi_cell(0, 5, f"Important: {text}")
        self.ln(3)
    
    def add_simple_table(self, headers, rows):
        col_count = len(headers)
        col_width = (self.w - self.l_margin - self.r_margin) / col_count
        
        # Header
        self.set_font("Helvetica", "B", 9)
        self.set_fill_color(59, 130, 246)
        self.set_text_color(255, 255, 255)
        for h in headers:
            self.cell(col_width, 7, h, border=1, fill=True, align="C")
        self.ln()
        
        # Rows
        self.set_font("Helvetica", "", 9)
        self.set_text_color(51, 51, 51)
        fill = False
        for row in rows:
            if fill:
                self.set_fill_color(241, 245, 249)
            else:
                self.set_fill_color(255, 255, 255)
            
            max_lines = 1
            for cell_text in row:
                lines = self.multi_cell(col_width, 6, cell_text, border=0, split_only=True)
                max_lines = max(max_lines, len(lines))
            
            row_height = max_lines * 6
            
            for i, cell_text in enumerate(row):
                x = self.get_x()
                y = self.get_y()
                self.rect(x, y, col_width, row_height, style="D")
                if fill:
                    self.set_fill_color(241, 245, 249)
                    self.rect(x + 0.3, y + 0.3, col_width - 0.6, row_height - 0.6, style="F")
                self.set_xy(x + 1, y + 1)
                self.multi_cell(col_width - 2, 6, cell_text)
                self.set_xy(x + col_width, y)
            self.ln(row_height)
            fill = not fill
        self.ln(4)


def create_pdf():
    pdf = PawLinkPDF()
    
    # ── Title Page ──
    pdf.add_page()
    pdf.ln(60)
    pdf.set_font("Helvetica", "B", 40)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 20, "PawLink", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 18)
    pdf.set_text_color(100, 116, 139)
    pdf.cell(0, 12, "Mobile App User Manual", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(0, 8, "Version 1.5.0  |  March 2026", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "Platform: Android (Expo / EAS)", align="C", new_x="LMARGIN", new_y="NEXT")
    
    # ── Table of Contents ──
    pdf.add_page()
    pdf.section_title("Table of Contents")
    toc_items = [
        "1. Introduction",
        "2. Getting Started",
        "3. Identity Verification",
        "4. Navigating the App",
        "5. Managing Your Pets",
        "6. Finding Breeding Matches",
        "7. Match Requests",
        "8. Chat & Messaging",
        "9. Breeding Contracts",
        "10. Shooter Workflow",
        "11. Subscriptions & Payments",
        "12. Activity & Notifications",
        "13. Profile & Settings",
        "14. Troubleshooting & FAQ",
        "15. Contact & Support",
    ]
    for item in toc_items:
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(0, 7, item, new_x="LMARGIN", new_y="NEXT")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 1
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("1. Introduction")
    pdf.body_text(
        "PawLink is a mobile pet breeding matchmaking platform that connects pet breeders with "
        "compatible breeding partners for their dogs and cats. The app also supports Shooters -- "
        "breeding handlers/facilitators who assist with the logistics of breeding arrangements."
    )
    
    pdf.sub3_title("Key Features")
    pdf.bullet("Smart Matching -- Swipe-based interface with compatibility scores.")
    pdf.bullet("AI-Powered Tools -- Breed identification, offspring prediction, OCR scanning.")
    pdf.bullet("Breeding Contracts -- In-app contract creation and management.")
    pdf.bullet("Secure Payments -- PayMongo with money pool escrow.")
    pdf.bullet("Real-Time Chat -- Direct messaging with read receipts.")
    pdf.bullet("Verification System -- Identity and document verification.")
    pdf.bullet("Vaccination Tracking -- Health record management.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 2
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("2. Getting Started")
    
    pdf.subsection_title("2.1 Installing the App")
    pdf.numbered_item(1, "Download the PawLink APK from the distribution link or QR code.")
    pdf.numbered_item(2, 'On Android, enable "Install from unknown sources" in Settings > Security.')
    pdf.numbered_item(3, "Open the APK file and tap Install.")
    pdf.numbered_item(4, "Once installed, tap Open to launch PawLink.")
    pdf.note_box("PawLink receives over-the-air (OTA) updates automatically.")
    
    pdf.subsection_title("2.2 Creating an Account")
    pdf.body_text("Tap Register on the login screen to begin the 4-step registration process:")
    
    pdf.sub4_title("Step 1 -- Account Setup")
    pdf.bullet("Enter email address (valid and unique).")
    pdf.bullet("Choose a username.")
    pdf.bullet("Create a password: 8+ chars, uppercase, number, special character.")
    
    pdf.sub4_title("Step 2 -- Personal Information")
    pdf.bullet("Enter first name and last name.")
    pdf.bullet("Select birthdate (minimum age: 13).")
    pdf.bullet("Choose sex (Male or Female).")
    
    pdf.sub4_title("Step 3 -- Address (Philippine format)")
    pdf.bullet("Street Address, Barangay, City/Municipality, Province, Postal Code (4 digits).")
    
    pdf.sub4_title("Step 4 -- Role Selection")
    pdf.bullet("Breeder -- Own pets, seeking breeding partners.")
    pdf.bullet("Shooter -- Facilitate breeding arrangements.")
    pdf.body_text("Select at least one role (can select both). Tap Submit to create your account.")
    
    pdf.subsection_title("2.3 Logging In")
    pdf.numbered_item(1, "Enter your email and password.")
    pdf.numbered_item(2, "Tap Sign In.")
    pdf.numbered_item(3, "You will be taken to the Home screen.")
    
    pdf.subsection_title("2.4 Forgot Password")
    pdf.body_text("Tap Forgot Password? on the login screen and follow instructions to reset via email.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 3
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("3. Identity Verification")
    pdf.body_text("Before adding pets or sending match requests, you must verify your identity.")
    
    pdf.subsection_title("3.1 Verifying Your Identity")
    pdf.body_text('Navigate to Profile > My Pets > "+" button, or Settings > Verification Status.')
    
    pdf.sub4_title("Step 1 -- Government ID Verification (Required)")
    pdf.bullet("Upload a photo of your government ID (camera or gallery).")
    pdf.bullet("OCR auto-fills: Full Name, ID Number, Issuing Authority, Dates.")
    pdf.bullet("Review and correct auto-filled info, then tap Next.")
    
    pdf.sub4_title("Step 2 -- Licensed Breeder Certificate (Optional)")
    pdf.bullet("Upload breeder's license if you have one. OCR auto-fills details.")
    pdf.bullet("Tap Skip if not applicable.")
    
    pdf.sub4_title("Step 3 -- Shooter Certificate (Optional)")
    pdf.bullet("Upload shooter certification if applicable.")
    pdf.bullet("Tap Skip if not applicable.")
    pdf.body_text("Tap Submit to send documents for review.")
    
    pdf.subsection_title("3.2 Adding Certificates Later")
    pdf.body_text("Go to Settings > Verification Status > Add Certificate on the relevant document type.")
    
    pdf.subsection_title("3.3 Verification Statuses")
    pdf.add_simple_table(
        ["Status", "Meaning"],
        [
            ["Not Submitted", "Document not yet uploaded."],
            ["Under Review", "Submitted, awaiting admin review."],
            ["Verified", "Document approved."],
            ["Rejected", "Rejected -- check reason and resubmit."],
        ]
    )
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 4
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("4. Navigating the App")
    
    pdf.subsection_title("4.1 Bottom Navigation Tabs")
    pdf.add_simple_table(
        ["Tab", "Description"],
        [
            ["Home", "Main dashboard -- top matches, nearby pets, shooters."],
            ["Matches", "Match requests, active matches, history."],
            ["Match (center)", "Full-screen swipe interface."],
            ["Activity", "Notification center."],
            ["Profile", "Dashboard, pets, settings."],
        ]
    )
    
    pdf.subsection_title("4.2 Switching Roles (Breeder / Shooter)")
    pdf.numbered_item(1, "Go to the Profile tab.")
    pdf.numbered_item(2, "Tap the Role Switcher toggle.")
    pdf.bullet("Breeder mode: Matching interface and pet discovery.")
    pdf.bullet("Shooter mode: Dashboard with offers and assignments.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 5
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("5. Managing Your Pets")
    
    pdf.subsection_title("5.1 Adding a New Pet")
    pdf.important_box("You must complete identity verification before adding a pet.")
    
    pdf.sub4_title("Step 1 -- Pet Information")
    pdf.bullet("Name, Species (Dog/Cat), Breed (AI auto-detection or manual), Sex, Birthdate, Height, Weight, Microchip ID.")
    
    pdf.sub4_title("Step 2 -- About Your Pet")
    pdf.bullet("Select behaviors (Friendly, Playful, Calm, etc.) and attributes. Write a description.")
    
    pdf.sub4_title("Step 3 -- Health Certificate")
    pdf.bullet("Upload health certificate photo. OCR auto-fills details.")
    
    pdf.sub4_title("Step 4 -- Photos")
    pdf.bullet("Upload minimum 3 photos. Tap to set primary photo.")
    
    pdf.sub4_title("Step 5 -- Preferences")
    pdf.bullet("Set preferred breeding partner criteria: breed, behaviors, attributes, age range.")
    
    pdf.subsection_title("5.2 Pet Profile Tabs")
    pdf.bullet("About -- Species, breed, sex, age, weight, height, behaviors, attributes, description.")
    pdf.bullet("Health -- Vaccination cards, shot records, document status.")
    pdf.bullet("Gallery -- Photo grid with primary photo.")
    pdf.bullet("History -- Breeding and litter history.")
    
    pdf.subsection_title("5.3 Pet Status Badges")
    pdf.add_simple_table(
        ["Badge", "Meaning"],
        [
            ["Available", "Ready for matching."],
            ["Pending", "Active request or awaiting verification."],
            ["Disabled", "Temporarily disabled."],
            ["Archived", "Not active."],
            ["Cooldown: Xd", "Post-breeding cooldown (X days)."],
        ]
    )
    
    pdf.subsection_title("5.4 Managing Vaccinations")
    pdf.numbered_item(1, "Open pet profile > Health tab.")
    pdf.numbered_item(2, "View dashboard: Total Cards, Verified, Pending, Overdue.")
    pdf.numbered_item(3, "Add shot: Tap vaccine card > Add Shot > enter date, upload photo > Submit.")
    pdf.numbered_item(4, "Manage protocols: Opt-in or change existing protocols.")
    
    pdf.subsection_title("5.5 Litters & Breeding History")
    pdf.bullet("View Total Litters, Total Pups/Kittens, Alive Count.")
    pdf.bullet("Tap a litter for: Overview, Offspring (allocation status), Health records.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 6
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("6. Finding Breeding Matches")
    
    pdf.subsection_title("6.1 The Home Screen (Breeder)")
    pdf.bullet("Top Matches -- Swipeable cards with photo, name, breed, compatibility score.")
    pdf.bullet("Nearby Pets -- Horizontal scroll of pets in your area.")
    pdf.bullet("Shooters -- Horizontal scroll of available shooters.")
    pdf.bullet("Browse Tabs -- Toggle between Pets and Shooters lists.")
    
    pdf.subsection_title("6.2 Swiping Through Matches")
    pdf.bullet("Swipe Right / Heart = Send match request.")
    pdf.bullet("Swipe Left / X = Pass (not shown again).")
    pdf.important_box("Only same species, opposite sex pets are shown. Your own pets are excluded.")
    
    pdf.subsection_title("6.3 Public Pet Profile")
    pdf.body_text("Tap any pet card to view 5 tabs: About, Health, Gallery, Litters, Compatibility.")
    pdf.body_text('"Send Match Request" button at the bottom.')
    
    pdf.subsection_title("6.4 Compatibility")
    pdf.bullet("Overall percentage score.")
    pdf.bullet("Trait breakdown: breed, health, age, location, behaviors, preferences.")
    
    pdf.subsection_title("6.5 AI Offspring Prediction")
    pdf.numbered_item(1, "From a pet profile, tap AI Offspring Prediction.")
    pdf.numbered_item(2, "View parent pets side by side.")
    pdf.numbered_item(3, "Tap Generate Prediction for AI-generated offspring image.")
    pdf.numbered_item(4, "View predicted breed mix, coat type, size, temperament.")
    pdf.note_box("AI generations per day limited by subscription tier (1/day on Free).")
    
    pdf.subsection_title("6.6 Search")
    pdf.numbered_item(1, "Tap search icon on Home.")
    pdf.numbered_item(2, "Filter by: All, Pets, Breeders, Shooters.")
    pdf.numbered_item(3, "For Pets: filter by Species and Sex.")
    pdf.numbered_item(4, "Tap a result to view their profile.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 7
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("7. Match Requests")
    
    pdf.subsection_title("7.1 Sending a Match Request")
    pdf.bullet("Swipe right on a card, tap Heart, or tap Send Match Request on a profile.")
    pdf.body_text("Requires: verified identity, active pet selected. Free tier may need per-request payment.")
    
    pdf.subsection_title("7.2 Incoming Requests")
    pdf.body_text("Matches tab > Requests > Incoming. Accept (Heart) or Decline (X).")
    pdf.body_text("Accepting creates a conversation automatically.")
    
    pdf.subsection_title("7.3 Outgoing Requests")
    pdf.body_text("Matches tab > Requests > Outgoing. Tap Cancel to withdraw.")
    
    pdf.subsection_title("7.4 Match History")
    pdf.body_text("Matches tab > History. Filter by: All, Declined, Cancelled.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 8
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("8. Chat & Messaging")
    
    pdf.subsection_title("8.1 Starting a Conversation")
    pdf.body_text('After a match is accepted, tap "Start Chat" in the Matches tab or tap the Activity notification.')
    
    pdf.subsection_title("8.2 Messaging Features")
    pdf.bullet("Real-time refresh every 5 seconds.")
    pdf.bullet("Your messages (right, blue) / theirs (left).")
    pdf.bullet("Timestamps and read receipts (double check marks).")
    pdf.bullet("Date headers between message groups.")
    pdf.bullet("Match info and contract card at the top.")
    pdf.bullet("Match timeline showing breeding progress.")
    
    pdf.subsection_title("8.3 Blocking & Reporting")
    pdf.numbered_item(1, "Tap the shield icon in the chat header.")
    pdf.numbered_item(2, "Block User (prevents communication) or Report User (select reason).")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 9
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("9. Breeding Contracts")
    
    pdf.subsection_title("9.1 Creating a Contract")
    pdf.body_text('In the conversation, tap "Create Contract". Complete 5 steps:')
    
    pdf.sub4_title("Step 1 -- Compensation Type")
    pdf.bullet("Money, Share Offspring, Goods/Foods (can combine).")
    
    pdf.sub4_title("Step 2 -- Compensation Details")
    pdf.bullet("Money: amount. Offspring: split method, value, picker. Goods: description.")
    
    pdf.sub4_title("Step 3 -- Shooter (Optional)")
    pdf.bullet("Shooter name, payment, location, conditions.")
    
    pdf.sub4_title("Step 4 -- Collateral & Timeline")
    pdf.bullet("Security deposit (50/50 split, held in escrow). Contract end date.")
    
    pdf.sub4_title("Step 5 -- Review & Submit")
    pdf.bullet("Review full summary and submit.")
    
    pdf.subsection_title("9.2 Accepting a Contract")
    pdf.numbered_item(1, "Open conversation > tap contract card.")
    pdf.numbered_item(2, "Review all terms.")
    pdf.numbered_item(3, "Accept or Reject.")
    pdf.numbered_item(4, "If accepted, both parties pay collateral.")
    
    pdf.subsection_title("9.3 Contract Lifecycle")
    pdf.add_simple_table(
        ["Stage", "Description"],
        [
            ["1. Create", "One party drafts terms."],
            ["2. Accept", "Other party reviews and accepts/rejects."],
            ["3. Pay Collateral", "Both pay 50% via PayMongo."],
            ["4. Submit Reports", "Breeding progress reports."],
            ["5. Mark Breeding", "Mark breeding as complete."],
            ["6. Record Offspring", "Record litter details."],
            ["7. Complete", "Finalize. Funds released."],
        ]
    )
    
    pdf.subsection_title("9.4 Contract Statuses")
    pdf.add_simple_table(
        ["Status", "Meaning"],
        [
            ["Draft", "Being created."],
            ["Pending", "Awaiting other party review."],
            ["Accepted", "Agreed -- awaiting collateral."],
            ["Shooter Requested", "Shooter invited."],
            ["Shooter Accepted", "Shooter accepted."],
            ["In Progress", "Breeding underway."],
            ["Breeding Complete", "Breeding confirmed."],
            ["Completed", "Finalized and closed."],
            ["Rejected / Cancelled", "Not proceeding."],
        ]
    )
    
    pdf.subsection_title("9.5 Contract Detail Tabs")
    pdf.bullet("Overview -- Terms, compensation, timeline, next action banner.")
    pdf.bullet("Payments -- Collateral, compensation, shooter payments.")
    pdf.bullet("Reports -- Breeding progress reports.")
    pdf.bullet("Breeding -- Log events, mark completion.")
    pdf.bullet("Offspring -- Litter details, individual offspring, allocation.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 10
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("10. Shooter Workflow")
    
    pdf.subsection_title("10.1 Shooter Home Dashboard")
    pdf.bullet("Stats: Active, Pending, Completed, Available + pet breakdown.")
    pdf.bullet("Current tab: Active/pending assignments.")
    pdf.bullet("Available tab: Open offers to accept.")
    pdf.bullet("Finished tab: Completed/failed with outcome badges.")
    
    pdf.subsection_title("10.2 Browsing & Accepting Offers")
    pdf.numbered_item(1, "Available tab > browse offers (both pets, payment, location).")
    pdf.numbered_item(2, "Tap offer for details.")
    pdf.numbered_item(3, 'Tap "Accept Offer" -- breeders notified for confirmation.')
    
    pdf.subsection_title("10.3 Active Assignments")
    pdf.numbered_item(1, "Current tab > tap assignment to join conversation.")
    pdf.numbered_item(2, "Assist with breeding logistics and submit reports.")
    pdf.numbered_item(3, "Payment processed through money pool on completion.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 11
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("11. Subscriptions & Payments")
    
    pdf.subsection_title("11.1 Subscription Plans")
    pdf.add_simple_table(
        ["Feature", "Free", "Standard", "Premium"],
        [
            ["Pets", "1", "More", "Unlimited"],
            ["Matches/Month", "3", "More", "Unlimited"],
            ["AI Predictions/Day", "1", "More", "Unlimited"],
            ["Billing", "--", "Monthly/Yearly", "Monthly/Yearly"],
        ]
    )
    
    pdf.subsection_title("11.2 Upgrading")
    pdf.numbered_item(1, "Profile > Settings > Subscription (or Home subscription badge).")
    pdf.numbered_item(2, "Browse plans, toggle Monthly/Yearly.")
    pdf.numbered_item(3, "Tap Subscribe > complete PayMongo payment.")
    
    pdf.subsection_title("11.3 Money Pool")
    pdf.body_text("Virtual escrow for breeding contract finances. Access via Profile > Settings > My Payments.")
    pdf.add_simple_table(
        ["Type", "Description"],
        [
            ["Held", "Funds in escrow for active contracts."],
            ["Frozen", "Frozen due to dispute/admin action."],
            ["Pending Deposits", "Collateral still processing."],
            ["Total Released", "Released from completed contracts."],
        ]
    )
    pdf.body_text("Transaction Tabs: Deposits, Releases, Refunds, Fee Deductions.")
    
    pdf.subsection_title("11.4 Filing a Dispute")
    pdf.body_text("My Payments > Disputes > File Dispute. Provide details and submit for admin review.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 12
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("12. Activity & Notifications")
    pdf.body_text("The Activity tab is your unified notification center with filter chips:")
    pdf.bullet("All, Verification, Matches, Messages, Shooter, Subscription, System.")
    pdf.ln(2)
    pdf.body_text("Verification notifications by severity:")
    pdf.bullet("Warnings -- Documents expiring soon.")
    pdf.bullet("Rejected -- Needs attention; tap Resubmit.")
    pdf.bullet("Under Review -- Submitted and pending.")
    pdf.bullet("Approved -- Verified and valid.")
    pdf.ln(2)
    pdf.body_text('Tap "Mark All as Read" to clear indicators. Tap any notification to navigate to its screen.')
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 13
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("13. Profile & Settings")
    
    pdf.subsection_title("13.1 Your Profile")
    pdf.bullet("Profile header: Photo, name, verification badge.")
    pdf.bullet("Role switcher toggle.")
    pdf.bullet("Dashboard: Current Breeding, Total Matches, Success Rate, Income.")
    pdf.bullet('My Pets: Grid with status badges. "+" to add new pet.')
    pdf.bullet("Settings: Quick access to account options.")
    
    pdf.subsection_title("13.2 Editing Your Profile")
    pdf.body_text("Profile > Settings > Account. Edit: Name, Contact Number, Birthdate, Sex, Profile Image. Email cannot be changed.")
    
    pdf.subsection_title("13.3 Privacy & Security")
    pdf.sub4_title("Change Password")
    pdf.body_text("Enter current password, new password (8+ chars, uppercase, number, special char), confirm. Tap Update Password.")
    
    pdf.sub4_title("Delete Account")
    pdf.body_text('Tap Delete Account > type "DELETE" > enter password > confirm. This is permanent.')
    
    pdf.subsection_title("13.4 Signing Out")
    pdf.body_text("Profile > Settings > Sign Out > Confirm.")
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 14
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("14. Troubleshooting & FAQ")
    
    faqs = [
        ("I can't add a pet.", "Complete identity verification first. Settings > Verification Status."),
        ("My verification was rejected.", "Check rejection reason in Verification Status and tap Resubmit."),
        ("No pets to match with.", "Ensure you have a pet registered and selected. Try Search."),
        ("Charged but subscription didn't activate.", "Return to app, check Settings > Subscription. Contact support if needed."),
        ("'Pending Shooter Request' shown.", "A shooter was invited but hasn't accepted yet. Wait for their response."),
        ("Accidentally passed on a pet.", "Use Search to find and send a match request from their profile."),
        ("Account suspended or banned.", "Contact pawlink.support@gmail.com."),
        ("How to switch Breeder/Shooter?", "Profile tab > role switcher toggle at the top."),
        ("How does the money pool work?", "Virtual escrow holding collateral and compensation. Released on completion, refunded on cancellation."),
        ("How many pets can I register?", "Free: 1 pet. Paid plans allow more. Check Subscription screen."),
    ]
    for q, a in faqs:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(0, 6, f"Q: {q}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5.5, f"A: {a}")
        pdf.ln(3)
    
    # ══════════════════════════════════════════════════════════════
    # SECTION 15
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf.section_title("15. Contact & Support")
    pdf.body_text("If you need assistance, have questions, or want to report an issue:")
    pdf.bullet("Email: pawlink.support@gmail.com")
    pdf.bullet("In-App: Use the Report feature in any conversation to flag concerns.")
    pdf.ln(20)
    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(0, 8, "PawLink -- Connecting Responsible Pet Breeders", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "© 2026 KHAT Development Team. All rights reserved.", align="C", new_x="LMARGIN", new_y="NEXT")
    
    # Save
    path = os.path.join(OUTPUT_DIR, "PawLink_User_Manual.pdf")
    pdf.output(path)
    print(f"PDF saved: {path}")
    return path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_docx()
    create_pdf()
    print("Done! All files generated.")
